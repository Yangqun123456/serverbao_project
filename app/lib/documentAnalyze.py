from app.lib.dataStructure import *
import jieba
import docx
import re
# 分类函数


def classify_function_points(doc_path):
    # 分类关键词
    ilf_keywords = ['文件', '数据库', '记录', '表', '数据', '界面', '子系统']
    eif_keywords = ['外部', '实体', '数据', '来源', '用户', '业务', '角色', '认证']
    ei_keywords = ['输入', '录入', '提供', '接受', '登录', '新增', '发起', '保存', '登记', '记录', '汇聚',
                   '生成', '优化', '构建', '识别', '整合', '采集', '归集', '研判', '投诉', '筛选', ]
    eo_keywords = ['输出', '显示', '生成', '展示', '总览']
    eq_keywords = ['查询', '搜索', '查找', '查阅', '对接', '接收', '链接',
                   '监管', '推送', '查看', '审核', '管理', '督办', '统计', '分析', '反馈']
    # 初始化
    ilf_sentences, eif_sentences, ei_sentences, eo_sentences, eq_sentences = [], [], [], [], []
    # 读取指定文档
    doc = docx.Document(doc_path)
    # 遍历文档中的段落
    for para in doc.paragraphs:
        text = para.text.strip()  # 去除文本首尾的空格
        if not text:  # 如果文本为空则跳过
            continue
        if is_start_with_number_and_dot_and_single_line(text):
            title = text
        if text != title:
            # 将文本进行分词
            words = list(jieba.cut(text))
            # 判断文本属于哪种功能点分类，并保存相关的句子
            if any(keyword in words for keyword in ilf_keywords):
                ilf_sentences.append(sentenceElement(text, title))
            elif any(keyword in words for keyword in eif_keywords):
                eif_sentences.append(sentenceElement(text, title))
            elif any(keyword in words for keyword in ei_keywords):
                ei_sentences.append(sentenceElement(text, title))
            elif any(keyword in words for keyword in eo_keywords):
                eo_sentences.append(sentenceElement(text, title))
            elif any(keyword in words for keyword in eq_keywords):
                eq_sentences.append(sentenceElement(text, title))

    return funPoints(ilf_sentences, eif_sentences, ei_sentences, eo_sentences, eq_sentences)


def is_start_with_number_and_dot_and_single_line(string):
    pattern = r"^\d+\..*\n?"
    if re.match(pattern, string):
        return True
    else:
        return False

# 计算工作量


def effortEstimates(funNumber, scaleFactor=1.39, developmentProductivity=7.04, operationsProductivity=0.85):
    development = developmentFactor()
    operations = operationsFactor()
    developmentEfforts = development.calculateEffort(
        funNumber*scaleFactor, developmentProductivity)
    operationsEfforts = operations.calculateEffort(
        funNumber*scaleFactor, operationsProductivity)
    return effortElement(developmentEfforts, operationsEfforts)

# 计算开发和运维成本


def costEstimates(effortElement, location):
    payload = paylaodPerson(location)
    developmentCost = effortElement.development*payload.development/10000
    operationsCost = effortElement.operations*payload.operations/10000
    return costElement(developmentCost, operationsCost)

# 更新DOC文档


def updateDocument(text, doc_path='./app/resource/需求文档.docx'):
    # 创建 Word 文档对象
    doc = docx.Document()
    # 将字符串分割为多个段落，并添加到 Word 文档对象中
    for paragraph_text in text.split('\r\n'):
        paragraph = doc.add_paragraph(paragraph_text)
    # 保存 Word 文档
    doc.save(doc_path)
